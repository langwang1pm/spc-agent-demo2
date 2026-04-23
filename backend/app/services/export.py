"""
导出服务 - 图表、数据、分析报告导出
"""
import io
import base64
from typing import Dict, Any, List
from datetime import datetime
import pandas as pd
from openpyxl import Workbook
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json


class ExportService:
    """导出服务"""
    
    @staticmethod
    def export_to_excel(data: List[List[Any]], 
                        headers: List[str] = None,
                        sheet_name: str = "Sheet1") -> bytes:
        """
        导出数据为Excel文件
        
        Args:
            data: 二维数组数据
            headers: 表头列表
            sheet_name: 工作表名称
        
        Returns:
            Excel文件字节数据
        """
        df = pd.DataFrame(data, columns=headers)
        
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name=sheet_name, index=False)
        
        return output.getvalue()
    
    @staticmethod
    def export_chart_data(data: Dict[str, Any]) -> bytes:
        """
        导出SPC图表数据为Excel
        
        Args:
            data: SPC计算结果
        
        Returns:
            Excel文件字节数据
        """
        wb = Workbook()
        ws = wb.active
        ws.title = "SPC数据"
        
        # 写入标题
        ws['A1'] = "SPC分析报告"
        ws['A1'].font = Pt(16)
        ws['A2'] = f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        
        # 写入统计结果
        stats = data.get("statistics", {})
        row = 4
        ws[f'A{row}'] = "统计指标"
        row += 1
        for key, value in stats.items():
            ws[f'A{row}'] = key
            ws[f'B{row}'] = value
            row += 1
        
        # 写入控制限
        limits = data.get("control_limits", {})
        row += 1
        ws[f'A{row}'] = "控制限"
        row += 1
        for key, value in limits.items():
            ws[f'A{row}'] = key
            ws[f'B{row}'] = value
            row += 1
        
        # 写入图表数据
        chart_data = data.get("chart_data", {})
        for chart_name, chart_info in chart_data.items():
            row += 2
            ws[f'A{row}'] = f"{chart_name}图表数据"
            ws[f'A{row}'].font = Pt(12)
            row += 1
            
            if "labels" in chart_info:
                ws[f'A{row}'] = "组别"
                ws[f'B{row}'] = "数值"
                if "ucl" in chart_info:
                    ws[f'C{row}'] = "UCL"
                if "cl" in chart_info:
                    ws[f'D{row}'] = "CL"
                if "lcl" in chart_info:
                    ws[f'E{row}'] = "LCL"
                row += 1
                
                for i, label in enumerate(chart_info["labels"]):
                    ws[f'A{row}'] = label
                    ws[f'B{row}'] = chart_info["data"][i]
                    if "ucl" in chart_info:
                        ws[f'C{row}'] = chart_info["ucl"]
                    if "cl" in chart_info:
                        ws[f'D{row}'] = chart_info["cl"]
                    if "lcl" in chart_info:
                        ws[f'E{row}'] = chart_info["lcl"]
                    row += 1
        
        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()
    
    @staticmethod
    def export_raw_data(data: List[List[float]], 
                       data_title: str = "原始数据") -> bytes:
        """
        导出原始数据为Excel
        
        Args:
            data: 二维数组数据
            data_title: 数据标题
        
        Returns:
            Excel文件字节数据
        """
        wb = Workbook()
        ws = wb.active
        ws.title = "原始数据"
        
        # 写入标题
        ws['A1'] = data_title
        ws['A1'].font = Pt(14)
        ws['A2'] = f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        
        # 写入数据
        row = 4
        ws[f'A{row}'] = "序号"
        ws[f'B{row}'] = "组号"
        ws[f'C{row}'] = "数据值"
        
        for i, group in enumerate(data):
            for j, value in enumerate(group):
                row += 1
                ws[f'A{row}'] = row - 4
                ws[f'B{row}'] = i + 1
                ws[f'C{row}'] = value
        
        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()
    
    @staticmethod
    def export_analysis_report(analysis_content: str,
                               data_info: Dict[str, Any],
                               report_format: str = "docx") -> bytes:
        """
        导出AI分析报告
        
        Args:
            analysis_content: 分析内容（Markdown格式）
            data_info: 数据信息
            report_format: 报告格式 (docx/html/pdf)
        
        Returns:
            报告文件字节数据
        """
        if report_format == "docx":
            return ExportService._export_to_docx(analysis_content, data_info)
        elif report_format == "html":
            return ExportService._export_to_html(analysis_content, data_info)
        elif report_format == "pdf":
            # PDF需要使用WeasyPrint，先导出HTML再转换
            html_content = ExportService._generate_html_report(analysis_content, data_info)
            return ExportService._html_to_pdf(html_content)
        else:
            raise ValueError(f"不支持的报告格式: {report_format}")
    
    @staticmethod
    def _export_to_docx(analysis_content: str, data_info: Dict[str, Any]) -> bytes:
        """导出为Word文档"""
        doc = Document()
        
        # 标题
        title = doc.add_heading('SPC智能分析报告', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 基本信息
        doc.add_heading('基本信息', level=1)
        doc.add_paragraph(f"数据标题: {data_info.get('data_title', 'N/A')}")
        doc.add_paragraph(f"数据源类型: {data_info.get('source_type', 'N/A')}")
        doc.add_paragraph(f"分析时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # 分析结果
        doc.add_heading('分析结果', level=1)
        
        # 简单解析Markdown内容并添加到Word
        lines = analysis_content.split('\n')
        for line in lines:
            if line.startswith('# '):
                doc.add_heading(line[2:], level=2)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=3)
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('| '):
                # 简单处理表格
                doc.add_paragraph(line)
            elif line.strip():
                doc.add_paragraph(line)
        
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
    
    @staticmethod
    def _export_to_html(analysis_content: str, data_info: Dict[str, Any]) -> bytes:
        """导出为HTML"""
        html_content = ExportService._generate_html_report(analysis_content, data_info)
        return html_content.encode('utf-8')
    
    @staticmethod
    def _generate_html_report(analysis_content: str, data_info: Dict[str, Any]) -> str:
        """生成HTML报告"""
        # 简单的Markdown到HTML转换
        html_content = analysis_content.replace('# ', '<h1>').replace('\n', '</h1>\n', 1)
        html_content = html_content.replace('\n## ', '</p><h2>').replace('\n# ', '</p><h1>')
        html_content = html_content.replace('\n', '<br>\n')
        
        return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>SPC分析报告</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; }}
        h1 {{ color: #333; border-bottom: 2px solid #1890ff; padding-bottom: 10px; }}
        h2 {{ color: #666; margin-top: 30px; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #1890ff; color: white; }}
    </style>
</head>
<body>
    <h1>SPC智能分析报告</h1>
    <p><strong>数据标题:</strong> {data_info.get('data_title', 'N/A')}</p>
    <p><strong>数据源类型:</strong> {data_info.get('source_type', 'N/A')}</p>
    <p><strong>分析时间:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    <hr>
    <div>{html_content}</div>
</body>
</html>"""
    
    @staticmethod
    def _html_to_pdf(html_content: str) -> bytes:
        """HTML转PDF（使用WeasyPrint）"""
        try:
            from weasyprint import HTML
            output = io.BytesIO()
            HTML(string=html_content).write_pdf(output)
            return output.getvalue()
        except Exception as e:
            raise Exception(f"PDF生成失败: {str(e)}，请确保已安装weasyprint")
    
    @staticmethod
    def export_chart_to_image(chart_data_url: str = None) -> str:
        """
        导出图表为图片（Base64）
        
        Note: 实际使用中，前端会将ECharts图表导出为图片，
        后端只需要接收和存储即可
        
        Returns:
            Base64编码的图片数据
        """
        # 这是一个占位函数
        return ""


# 全局实例
export_service = ExportService()


def export_excel(data: List[List[Any]], headers: List[str] = None) -> bytes:
    """导出Excel便捷函数"""
    return export_service.export_to_excel(data, headers)


def export_spc_chart(data: Dict[str, Any]) -> bytes:
    """导出SPC图表数据"""
    return export_service.export_chart_data(data)


def export_raw(data: List[List[float]], title: str = "原始数据") -> bytes:
    """导出原始数据"""
    return export_service.export_raw_data(data, title)


def export_report(content: str, data_info: Dict[str, Any], 
                  format: str = "docx") -> bytes:
    """导出分析报告"""
    return export_service.export_analysis_report(content, data_info, format)